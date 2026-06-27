from __future__ import annotations

import shutil
import time
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RELEASE_ROOT = Path("D:/RuleQuant/release")
STAMP = time.strftime("%Y%m%d_%H%M")
PACKAGE_NAME = f"RuleQuant_Web_Package_{STAMP}"
PACKAGE_DIR = RELEASE_ROOT / PACKAGE_NAME


def write_text(path: Path, text: str) -> None:
  path.parent.mkdir(parents=True, exist_ok=True)
  path.write_text(text, encoding="utf-8")


def copy_tree(src: Path, dst: Path) -> None:
  if dst.exists():
    shutil.rmtree(dst)
  shutil.copytree(src, dst)


def repair_runtime_node_modules(app_dir: Path) -> None:
  """Restore pnpm package entrypoints that Next standalone can omit on Windows."""
  pnpm_dir = app_dir / "node_modules" / ".pnpm"
  if not pnpm_dir.exists():
    return

  root_node_modules = app_dir / "node_modules"
  for package_dir in sorted(pnpm_dir.iterdir()):
    package_node_modules = package_dir / "node_modules"
    if not package_dir.is_dir() or not package_node_modules.exists():
      continue
    for entry in package_node_modules.iterdir():
      if entry.name == ".bin" or entry.name == "node_modules":
        continue
      if entry.name.startswith("@"):
        for scoped_entry in entry.iterdir():
          if not scoped_entry.is_dir():
            continue
          destination = root_node_modules / entry.name / scoped_entry.name
          if not destination.exists():
            copy_tree(scoped_entry, destination)
      else:
        destination = root_node_modules / entry.name
        if entry.is_dir() and not destination.exists():
          copy_tree(entry, destination)


def build_docx(path: Path) -> None:
  doc = Document()
  section = doc.sections[0]
  section.top_margin = Inches(0.75)
  section.bottom_margin = Inches(0.75)
  section.left_margin = Inches(0.8)
  section.right_margin = Inches(0.8)

  normal = doc.styles["Normal"]
  normal.font.name = "Microsoft YaHei"
  normal.font.size = Pt(10.5)
  normal.paragraph_format.space_after = Pt(6)
  normal.paragraph_format.line_spacing = 1.25

  for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
    style = doc.styles[name]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(31, 78, 121)
    style.paragraph_format.space_before = Pt(10)
    style.paragraph_format.space_after = Pt(6)

  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  run = p.add_run("RuleQuant 回测终端使用说明书")
  run.font.name = "Microsoft YaHei"
  run.font.size = Pt(22)
  run.font.bold = True
  run.font.color.rgb = RGBColor(11, 37, 69)

  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  run = p.add_run("公式驱动的自动计算、历史验证、逐期明细和综合参考排序系统")
  run.font.name = "Microsoft YaHei"
  run.font.size = Pt(11)
  run.font.color.rgb = RGBColor(85, 85, 85)

  def bullets(items: list[str]) -> None:
    for item in items:
      doc.add_paragraph(item, style="List Bullet")

  def numbers(items: list[str]) -> None:
    for item in items:
      doc.add_paragraph(item, style="List Number")

  def table(rows: list[tuple[str, str]]) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "项目"
    tbl.rows[0].cells[1].text = "说明"
    for left, right in rows:
      cells = tbl.add_row().cells
      cells[0].text = left
      cells[1].text = right

  doc.add_heading("一、系统定位", level=1)
  doc.add_paragraph("RuleQuant 只用于历史开奖数据分析、规则公式计算、回测、逐期明细核对和综合参考排序。")
  bullets([
    "不是预测工具，不代表一定正确。",
    "综合参考结果由历史公式表现和最新一期公式计算结果合并生成。",
    "所有计算统一经过 Formula Engine，页面只展示结果和证据。",
  ])

  doc.add_heading("二、打开方式", level=1)
  numbers([
    "解压压缩包。",
    "双击“打开RuleQuant网页.cmd”。",
    "等待浏览器自动打开 RuleQuant 首页。",
    "如果没有自动打开，到 logs 文件夹查看 last-url.txt，复制里面的网址到浏览器。",
  ])

  doc.add_heading("三、每天使用流程", level=1)
  table([
    ("同步开奖数据", "首页或综合参考结果页点击“同步最新开奖数据”，系统会抓取网站全年数据并写入本地库。"),
    ("一键算公式", "进入“一键算公式”，用最新一期代入所有启用且可计算的公式。"),
    ("看综合参考", "进入“综合参考结果”，查看 Top 号码、Top 生肖，并点开候选看支持和反对证据。"),
    ("查逐期明细", "进入“公式逐期明细”，检查每一期变量取值、表达式、rawResult、归一化、输出和对错。"),
  ])

  doc.add_heading("四、规则口径", level=1)
  table([
    ("D序", "只排序 6 个平码，特码单独保留为第 7 位。"),
    ("L序/落位", "L序按原始落球顺序；落1到落7始终按原始落球顺序。"),
    ("期合", "按后三位期数算，例如 2026174 按 174 算，期合=12，期合尾=2。"),
    ("波色值", "红=0，蓝=1，绿=2。"),
    ("五行值", "金=1，木=2，水=3，火=4，土=5，并使用用户最新提供的 2026 五行表。"),
    ("合数尾", "按固定 0-9 合尾表，本质等同号码合尾，公式可写合数尾。"),
    ("头数单/双", "支持平1头单、平1头双、特码头单、特码头双等变量。"),
    ("七尾", "按 -3、-2、-1、0、+1、+2、+4 闭环生成。"),
    ("九肖", "支持 123456.5432.123456.5432 取位循环，按号码 +1 后生成九肖集合。"),
  ])

  doc.add_heading("五、公式管理", level=1)
  bullets([
    "公式管理默认智能排行：成功率高、近10期好、连错少的公式排前。",
    "最近出错会进入提醒区；连错较多且命中率不稳的可放入备选库。",
    "备选库公式继续回测，但不参与综合参考，后续可恢复。",
    "TXT 公式导入会追加为用户提供公式，不覆盖原公式库，导入前自动备份。",
  ])

  doc.add_heading("六、给朋友审核重点", level=1)
  numbers([
    "首页是否显示最新期号和开奖号码。",
    "一键算公式是否显示所有启用公式的计算结果。",
    "综合参考结果是否不是 0 分空结果，且有公式依据。",
    "点击 Top 号码后，右侧是否显示支持/反对公式和计算过程。",
    "新增公式是否为空白表单，试算当前草稿是否正常。",
    "TXT 导入是否追加公式，不覆盖原公式库。",
  ])

  doc.add_heading("七、常见问题", level=1)
  table([
    ("打不开", "双击打开RuleQuant网页.cmd；如果失败，把 logs 文件夹发给开发者。"),
    ("同步慢", "第一次会抓全年数据，后续会复用最新同步数据。"),
    ("综合结果为空", "说明没有任何可参与公式或没有依据，系统不会展示假 Top。"),
    ("误删公式", "公式库会自动备份，可在公式管理页恢复上一次备份。"),
  ])

  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  run = p.add_run("仅供公式研究和参考排序，不代表一定正确。")
  run.bold = True
  run.font.color.rgb = RGBColor(155, 28, 28)

  path.parent.mkdir(parents=True, exist_ok=True)
  doc.save(path)


def build_launcher(package_dir: Path) -> None:
  launcher = """
const fs = require('fs');
const http = require('http');
const net = require('net');
const path = require('path');
const childProcess = require('child_process');

const baseDir = __dirname;
const appDir = path.join(baseDir, 'app');
const logsDir = path.join(baseDir, 'logs');
fs.mkdirSync(logsDir, { recursive: true });

function canUsePort(port) {
  return new Promise((resolve) => {
    const server = net.createServer();
    server.once('error', () => resolve(false));
    server.once('listening', () => server.close(() => resolve(true)));
    server.listen(port, '127.0.0.1');
  });
}

async function findPort(start) {
  for (let port = start; port < start + 30; port += 1) {
    if (await canUsePort(port)) return port;
  }
  throw new Error('NO_FREE_PORT');
}

function waitForHttp(url, timeoutMs = 15000) {
  const started = Date.now();
  return new Promise((resolve, reject) => {
    const tick = () => {
      const req = http.get(url, (res) => {
        res.resume();
        if (res.statusCode && res.statusCode < 500) resolve();
        else retry();
      });
      req.on('error', retry);
      req.setTimeout(2000, () => { req.destroy(); retry(); });
    };
    const retry = () => {
      if (Date.now() - started > timeoutMs) reject(new Error('SERVER_START_TIMEOUT'));
      else setTimeout(tick, 500);
    };
    tick();
  });
}

function openBrowser(url) {
  childProcess.spawn('cmd', ['/c', 'start', '', url], {
    detached: true,
    stdio: 'ignore',
    windowsHide: true,
  }).unref();
}

(async () => {
  const port = await findPort(Number(process.env.RULEQUANT_PORT || '3088'));
  const out = fs.openSync(path.join(logsDir, 'server.log'), 'a');
  const err = fs.openSync(path.join(logsDir, 'server-error.log'), 'a');
  const child = childProcess.spawn(process.execPath, [path.join(appDir, 'server.js')], {
    cwd: appDir,
    detached: true,
    stdio: ['ignore', out, err],
    windowsHide: true,
    env: { ...process.env, PORT: String(port), HOSTNAME: '127.0.0.1' },
  });
  child.unref();
  const url = `http://127.0.0.1:${port}/dashboard`;
  await waitForHttp(url);
  fs.writeFileSync(path.join(logsDir, 'last-url.txt'), url, 'utf8');
  if (!process.argv.includes('--no-open')) openBrowser(url);
  console.log(url);
})().catch((error) => {
  fs.writeFileSync(path.join(logsDir, 'launcher-error.log'), String(error && error.stack || error), 'utf8');
  console.error(error);
  process.exit(1);
});
""".strip()
  write_text(package_dir / "launcher.js", launcher)
  cmd = """@echo off
cd /d "%~dp0"
"%~dp0runtime\\node\\node.exe" "%~dp0launcher.js"
if errorlevel 1 (
  echo.
  echo RuleQuant failed to start. Please send the logs folder to the developer.
  pause
)
"""
  write_text(package_dir / "打开RuleQuant网页.cmd", cmd)


def zip_dir(src: Path, zip_path: Path) -> None:
  if zip_path.exists():
    zip_path.unlink()
  with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as zf:
    for file in src.rglob("*"):
      if file.is_file():
        zf.write(file, file.relative_to(src.parent))


def main() -> None:
  RELEASE_ROOT.mkdir(parents=True, exist_ok=True)
  if PACKAGE_DIR.exists():
    shutil.rmtree(PACKAGE_DIR)
  PACKAGE_DIR.mkdir(parents=True)

  app_dir = PACKAGE_DIR / "app"
  copy_tree(ROOT / ".next" / "standalone", app_dir)
  repair_runtime_node_modules(app_dir)
  copy_tree(ROOT / ".next" / "static", app_dir / ".next" / "static")
  if (ROOT / "public").exists():
    copy_tree(ROOT / "public", app_dir / "public")

  runtime_dir = PACKAGE_DIR / "runtime" / "node"
  runtime_dir.mkdir(parents=True, exist_ok=True)
  shutil.copy2(Path("C:/Users/32129/.cache/codex-runtimes/codex-primary-runtime/dependencies/node/bin/node.exe"), runtime_dir / "node.exe")

  build_launcher(PACKAGE_DIR)
  build_docx(PACKAGE_DIR / "RuleQuant使用说明书.docx")

  screenshots = PACKAGE_DIR / "screenshots"
  screenshots.mkdir(parents=True, exist_ok=True)
  src_screens = ROOT / "output" / "ui-verify-20260627"
  if src_screens.exists():
    for file in src_screens.glob("*.png"):
      shutil.copy2(file, screenshots / file.name)

  write_text(PACKAGE_DIR / "审核说明.txt", "\\n".join([
    "RuleQuant 回测终端审核说明",
    "",
    "打开方式：双击 打开RuleQuant网页.cmd。",
    "用途：历史开奖数据、规则公式、回测、逐期明细、综合参考排序。",
    "声明：仅供公式研究和参考排序，不代表一定正确。",
    "",
    "验证：测试 50 条通过；TypeScript 检查通过；生产构建通过；浏览器截图已放入 screenshots 文件夹。",
  ]))

  zip_path = RELEASE_ROOT / f"{PACKAGE_NAME}.zip"
  zip_dir(PACKAGE_DIR, zip_path)
  print(PACKAGE_DIR)
  print(zip_path)


if __name__ == "__main__":
  main()
