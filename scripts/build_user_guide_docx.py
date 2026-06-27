from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_font(run, size=None, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_font(run, size=9.5, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header = table.rows[0].cells
    for index, text in enumerate(rows[0]):
        set_cell_shading(header[index], "E8EEF5")
        set_cell_text(header[index], text, bold=True, color="1F4D78")

    for row in rows[1:]:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            set_cell_text(cells[index], str(text))

    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)

    doc.add_paragraph("")
    return table


def add_paragraphs(doc, paragraphs):
    for text in paragraphs:
        doc.add_paragraph(text)


def build_docx(output_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in [
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RuleQuant 回测终端使用说明书")
    set_font(run, size=22, bold=True, color="0F766E")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("可转发试用版 | 2026-06-26")
    set_font(run, size=10, color="64748B")

    doc.add_heading("一、最快打开方式", level=1)
    add_paragraphs(
        doc,
        [
            "解压压缩包后，优先双击“打开RuleQuant网页.hta”。它会显示一个网页式启动页，点击“启动并打开 RuleQuant”即可打开浏览器。",
            "如果电脑禁止 HTA，可以双击“启动RuleQuant服务.bat”，然后在浏览器访问 http://127.0.0.1:3030/dashboard。",
            "压缩包内已经带 node.exe 和运行依赖，普通用户不需要自己安装 Node.js。",
        ],
    )
    add_table(
        doc,
        [
            ["文件", "作用"],
            ["打开RuleQuant网页.hta", "网页式启动器，点击按钮启动本地服务并打开系统。"],
            ["打开RuleQuant.html", "普通网页入口，服务启动后可点击进入系统。"],
            ["启动RuleQuant服务.bat", "后备启动方式，电脑拦截 HTA 时使用。"],
            ["RuleQuant使用说明书.docx", "本说明书。"],
        ],
        [2.1, 4.4],
    )

    doc.add_heading("二、每天同步和计算流程", level=1)
    add_paragraphs(
        doc,
        [
            "1. 第一步：打开首页，系统会自动同步网站全年开奖数据。看到最新期号和最新开奖号码后再继续。",
            "2. 第二步：点击“一键计算全部公式”，系统会把最新一期代入所有已启用且可计算的用户提供公式。",
            "3. 第三步：进入“综合参考结果”，查看参考号码、参考生肖，并点击候选号码查看支持和反对证据。",
            "4. 第四步：进入“公式逐期明细”，检查每条公式每一期怎么算、哪期对、哪期错。最新一期没有下一期开奖时会显示“待验证”。",
        ],
    )

    doc.add_heading("三、核心页面说明", level=1)
    add_table(
        doc,
        [
            ["页面", "普通用户要看什么"],
            ["首页", "三步式流程：同步数据、一键算公式、查看综合参考。顶部必须显示最新期。"],
            ["开奖数据", "每天新增的期开奖数据会排在第一行。D序只排序 6 个平码，特码单独保留。"],
            ["一键算公式", "查看每条公式本期输出结果，包含变量取值、表达式、原始结果和最终输出。"],
            ["公式逐期明细", "每一期完整流水账：当前开奖、变量、计算过程、结果映射、下一期开奖、正确/错误/待验证。"],
            ["综合参考结果", "把历史表现和最新一期公式输出合并，生成参考号码和生肖；点击候选能看依据。"],
            ["公式管理", "启用、停用、编辑、删除、备份、恢复、放入备选库。"],
            ["公式提醒 / 备选库", "错了会提醒；连错太多且命中率不稳的公式可以放入备选库，继续回测但暂不参与综合参考。"],
            ["手动公式组合", "在综合参考页手动勾选几条公式，单独查看这组公式合并后的号码和生肖。"],
        ],
        [1.55, 4.95],
    )

    doc.add_heading("四、新增规则口径", level=1)
    add_paragraphs(
        doc,
        [
            "D序：只把 6 个平码从小到大排序，特码永远单独保留为第 7 位，不参与排序。",
            "期合 / 期合尾：按后三位期数计算，例如 2026174 的期数是 174，期合是 1+7+4=12，期合尾是 2。",
            "单双自用 4455：按 4、4、5、5 循环取位；取平4时加 1 后看单双，取平5时加 2 后看单双。",
            "波色、大小、单双：系统已作为号码属性参与公式计算和综合参考证据。",
            "样例核对：用于检查程序有没有理解错，不是用户提供公式参与综合参考的强制门槛。",
        ],
    )

    doc.add_heading("五、综合参考结果怎么理解", level=1)
    add_paragraphs(
        doc,
        [
            "综合参考结果不是固定历史排名，也不是保证结果。",
            "历史数据用于判断公式过去表现；最新开奖用于计算本期公式输出；综合结果把所有公式的支持和排除合并成参考排序。",
            "如果没有同步新开奖，也没有修改公式，结果应基本不变；同步新期开奖、修改公式、修改生肖/波色/五行表后必须重新计算。",
            "页面不得理解为必中、稳赚、包准、精准预测或投注建议，只用于公式研究和参考排序。",
        ],
    )

    doc.add_heading("六、发给别人前检查", level=1)
    add_table(
        doc,
        [
            ["检查项", "通过标准"],
            ["能打开", "双击“打开RuleQuant网页.hta”，点击按钮后浏览器出现首页。"],
            ["数据最新", "首页和开奖数据页显示最新期号、最新开奖号码、同步记录数。"],
            ["公式能算", "一键算公式页面有每条公式输出，不是空白。"],
            ["明细完整", "公式逐期明细能看到最新期，最新期应是“待验证”或已有下一期判断。"],
            ["综合不为 0", "综合参考结果显示参与公式数、生成证据数，候选号码能点开看依据。"],
            ["备份可用", "公式管理页可导出公式库 JSON，修改前会自动备份。"],
        ],
        [1.8, 4.7],
    )

    doc.add_heading("七、问题反馈时请发这些信息", level=1)
    add_paragraphs(
        doc,
        [
            "截图：问题页面截图，最好带顶部最新期号。",
            "操作：刚刚点了哪个按钮，输入了什么号码或公式。",
            "数据：当前最新期号、开奖号码、是否同步成功。",
            "公式：公式名称、公式逐期明细中出错的期号。",
        ],
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("RuleQuant 回测终端 | 仅供公式研究和参考排序")
    set_font(run, size=9, color="64748B")

    doc.save(output_path)


if __name__ == "__main__":
    release_dir = Path("D:/RuleQuant/release")
    release_dir.mkdir(parents=True, exist_ok=True)
    build_docx(release_dir / "RuleQuant使用说明书_20260626_可转发版.docx")
